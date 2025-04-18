from docx import Document
import os
from django.conf import settings
from .generation import generate_worksheet_content
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, Spacer
from reportlab.platypus import SimpleDocTemplate
def create_sheet(sheet):
    sheet.content = generate_worksheet_content(sheet)
    create_worksheet_files(sheet)
    sheet.save()

def update_worksheet_files(sheet):
    create_worksheet_files(sheet)

def create_worksheet_files(sheet):
    # Setup paths
    media_dir = os.path.join('worksheets', str(sheet.user.id))
    os.makedirs(os.path.join(settings.MEDIA_ROOT, media_dir), exist_ok=True)

    filename = f"worksheet_{sheet.id}_{sheet.created_at.strftime('%Y%m%d_%H%M%S')}"
    docx_path = os.path.join(media_dir, f"{filename}.docx")
    pdf_path = os.path.join(media_dir, f"{filename}.pdf")
    
    # Create files with absolute paths
    create_docx(sheet, os.path.join(settings.MEDIA_ROOT, docx_path))
    create_pdf(sheet, os.path.join(settings.MEDIA_ROOT, pdf_path))

    # Save relative paths to the model
    sheet.docx_file.name = docx_path
    sheet.pdf_file.name = pdf_path
    sheet.save()
    

def create_docx(sheet, filepath):
    # Create DOCX file
    doc = Document()
    doc.add_paragraph('Name: ___________________________ Date: ___________________________')
    doc.add_paragraph(sheet.content)
    doc.save(filepath)


def create_pdf(sheet, filepath):
    pdf_doc = SimpleDocTemplate(
                filepath,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

    # Create the PDF content
    styles = getSampleStyleSheet()
    story = []

    # Add name and date
    story.append(Paragraph('Name: ___________________________ Date: ___________________________', styles['Normal']))
    story.append(Spacer(1, 12))

    # Split content into paragraphs and add them
    paragraphs = sheet.content.split('\n')
    for para in paragraphs:
        if para.strip():  # Skip empty lines
            story.append(Paragraph(para, styles['Normal']))
            story.append(Spacer(1, 12))

    # Build the PDF
    pdf_doc.build(story)
    